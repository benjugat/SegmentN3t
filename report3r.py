from docx import *
from docx.shared import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

import sys,os
try:
	import argparse
except:
	print('[!] argparse is not installed. Try "pip install argparse"')
	sys.exit(0)

def check_pass(data, paragraph):

	fail_text = 'FAIL ' + u'\u2718'
	pass_text = 'PASS ' + u'\u2713'
	
	for i in data:
		if "open" in i:
			#FAIL
			paragraph_run = paragraph.add_run(fail_text)
			ffont = paragraph_run.font
			ffont.color.rgb =  RGBColor(0x96, 0x00, 0x00)
			paragraph_run.bold=True
			return
	#PASS
	paragraph_run = paragraph.add_run(pass_text)
	ffont = paragraph_run.font
	ffont.color.rgb =  RGBColor(0x00, 0x96, 0x00)
	paragraph_run.bold=True


def make_table(route, name, doc):

	# Table  
	table = doc.add_table(rows=3, cols=1, style='Tabla')
	table.autofit = True
	table.allow_autofit = True

	hdr_cells = table.rows[0].cells
	hdr_cells[0].text = 'Resultado de las pruebas ' + name
	hdr_cells[0].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
	hdr_cells = table.rows[1].cells
	hdr_cells[0].text = 'Técnicas de Evasión'
	hdr_cells[0].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
	hdr_cells = table.rows[2].cells



	# TABLE

	# NORMAL
	f = open(route + '/normal.nmap', 'r')
	data = f.read().split('\n')
	paragraph = hdr_cells[0].add_paragraph()
	paragraph_run = paragraph.add_run('Port Scan (Normal):\t')
	paragraph_run.bold=True
	paragraph_run.underline=True
	check_pass(data, paragraph)

	paragraph = hdr_cells[0].add_paragraph('No se utiliza ningún tipo de evasión.\n')
	paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

	for i in data:
		hdr_cells[0].add_paragraph(i)
	hdr_cells[0].add_paragraph('')

	# ICMP SCAN
	f = open(route + '/hping.txt', 'r')
	data = f.read().split('\n')
	paragraph = hdr_cells[0].add_paragraph()
	paragraph_run = paragraph.add_run('ICMP Scan:\t')
	paragraph_run.bold=True
	paragraph_run.underline=True
	check_pass(data, paragraph)

	paragraph = hdr_cells[0].add_paragraph('Se utiliza el sondeo a través de ICMP para verificar interconexión.\n')
	paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
	
	for i in data:
		hdr_cells[0].add_paragraph(i)
	hdr_cells[0].add_paragraph('')

	# DECOY
	f = open(route + '/decoy.nmap', 'r')
	data = f.read().split('\n')
	paragraph = hdr_cells[0].add_paragraph()
	paragraph_run = paragraph.add_run('Decoy:\t')
	paragraph_run.bold=True
	paragraph_run.underline=True
	check_pass(data, paragraph)

	paragraph = hdr_cells[0].add_paragraph('Se falsifican paquetes de otros hosts.\n')
	paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
	
	for i in data:
		hdr_cells[0].add_paragraph(i)
	hdr_cells[0].add_paragraph('')

	# FRAG
	f = open(route + '/frag.nmap', 'r')
	data = f.read().split('\n')
	paragraph = hdr_cells[0].add_paragraph()
	paragraph_run = paragraph.add_run('Fragmentación de Paquetes:\t')
	paragraph_run.bold=True
	paragraph_run.underline=True
	check_pass(data, paragraph)
	paragraph = hdr_cells[0].add_paragraph('Se fragmentan los paquetes enviados.\n')
	paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
	
	for i in data:
		hdr_cells[0].add_paragraph(i)
	hdr_cells[0].add_paragraph('')

	# MTU
	f = open(route + '/mtu16.nmap', 'r')
	data = f.read().split('\n')
	paragraph = hdr_cells[0].add_paragraph()
	paragraph_run = paragraph.add_run('MTU:\t')
	paragraph_run.bold=True
	paragraph_run.underline=True
	check_pass(data, paragraph)
	paragraph = hdr_cells[0].add_paragraph('Se establece la unidad de transmisión máxima específica para cada paquete.\n')
	paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
	
	for i in data:
		hdr_cells[0].add_paragraph(i)
	hdr_cells[0].add_paragraph('')

	# BADSUM
	f = open(route + '/badsum.nmap', 'r')
	data = f.read().split('\n')
	paragraph = hdr_cells[0].add_paragraph()
	paragraph_run = paragraph.add_run('Bad Checksums:\t')
	paragraph_run.bold=True
	paragraph_run.underline=True
	check_pass(data, paragraph)
	paragraph = hdr_cells[0].add_paragraph('Se utiliza una suma de comprobación errónea.\n')
	paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
	
	for i in data:
		hdr_cells[0].add_paragraph(i)
	hdr_cells[0].add_paragraph('')

	# SOURCE
	f53 = open(route + '/source53.nmap', 'r')
	data53 = f53.read().split('\n')
	f80 = open(route + '/source80.nmap', 'r')
	data80 = f80.read().split('\n')
	f88 = open(route + '/source88.nmap', 'r')
	data88 = f88.read().split('\n')
	f443 = open(route + '/source443.nmap', 'r')
	data443 = f443.read().split('\n')
	paragraph = hdr_cells[0].add_paragraph()
	paragraph_run = paragraph.add_run('Distinto puerto origen:\t')
	paragraph_run.bold=True
	paragraph_run.underline=True
	check_pass(data53+data80+data88+data443, paragraph)
	paragraph = hdr_cells[0].add_paragraph('Se cambia el puerto de origen desde el cual se realizan los escaneos.\n')
	paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
	
	## SOURCE 53
	hdr_cells[0].add_paragraph('Prueba con puerto de origen 53 (DNS):\n')
	paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
	
	for i in data53:
		hdr_cells[0].add_paragraph(i)
	hdr_cells[0].add_paragraph('')

	## SOURCE 80
	paragraph = hdr_cells[0].add_paragraph('Prueba con puerto origen 80 (HTTP):\n')
	paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
	for i in data80:
		hdr_cells[0].add_paragraph(i)
	hdr_cells[0].add_paragraph('')

	## SOURCE 88
	paragraph = hdr_cells[0].add_paragraph('Prueba con puerto origen 88 (Kerberos):\n')
	paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
	for i in data88:
		hdr_cells[0].add_paragraph(i)
	hdr_cells[0].add_paragraph('')

	## SOURCE 443
	paragraph = hdr_cells[0].add_paragraph('Prueba con puerto origen 443 (HTTPS):\n')
	paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
	for i in data443:
		hdr_cells[0].add_paragraph(i)
	hdr_cells[0].add_paragraph('')

	# FLAG
	fnull = open(route + '/flag_null.nmap', 'r')
	datanull = fnull.read().split('\n')
	fxmas = open(route + '/flag_xmas.nmap', 'r')
	dataxmas = fxmas.read().split('\n')
	ffin = open(route + '/flag_fin.nmap', 'r')
	datafin = ffin.read().split('\n')

	paragraph = hdr_cells[0].add_paragraph()
	paragraph_run = paragraph.add_run('Stateless:\t')
	paragraph_run.bold=True
	paragraph_run.underline=True
	check_pass(datanull + dataxmas + datafin, paragraph)
	paragraph = hdr_cells[0].add_paragraph('Se utilizan distintos tipos de flag como NULL, XMAS y FIN con la finalidad de evadir firewalls stateless.\n')
	paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

	## FLAG NULL
	paragraph = hdr_cells[0].add_paragraph('Prueba con flag NULL activado:\n')
	paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
	
	for i in datanull:
		hdr_cells[0].add_paragraph(i)
	hdr_cells[0].add_paragraph('')

	## FLAG XMAS
	paragraph = hdr_cells[0].add_paragraph('Prueba con flag XMAS activado:\n')
	paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
	for i in dataxmas:
		hdr_cells[0].add_paragraph(i)
	hdr_cells[0].add_paragraph('')

	## FLAG FIN
	paragraph = hdr_cells[0].add_paragraph('Prueba con flag FIN activado:\n')
	paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
	for i in datafin:
		hdr_cells[0].add_paragraph(i)
	hdr_cells[0].add_paragraph('')


def main():

	# Parsing arguments
	parser = argparse.ArgumentParser(description='Report3r is used for making reports from SegmentN3t tool.\n\t\t\n Example: $ python3 report3r.py -r /tmp/report/2050-09-21_08:48:46/', epilog='Thanks for using me!')
	parser.add_argument('-r', '--route', action='store', dest='route', help='Route of results')
	global args
	args =  parser.parse_args()

	if len(sys.argv) < 2:
		parser.print_help()
		sys.exit(0)

	# create an instance of a word document
	doc = Document("template.docx")

	doc.add_heading('Results of %s' % (args.route), 1)
	doc_paragraph = doc.add_paragraph('')

	# discovering files
	for network_dir in os.listdir(args.route):
		for subnetwork_dir in os.listdir(os.path.join(args.route, network_dir)):
			make_table(os.path.join(args.route,network_dir,subnetwork_dir), subnetwork_dir, doc)
			doc.add_paragraph('')
			print("[+] Making table of results in : %s" % (os.path.join(args.route,network_dir,subnetwork_dir)))

	# saving the document
	doc.save("./report.docx")


try:
	if __name__ == "__main__":
		main()
except KeyboardInterrupt:
	print("[!] Keyboard Interrupt. Shutting down")
